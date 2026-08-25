from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = Path(r"C:\Users\wuaomo\Documents\AI Web Form Agent\Regional_AI_Digital_Transformation_Intern_Interview_Prep.docx")


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "1F2937"
MUTED = "4B5563"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
CALLOUT = "F4F6F9"
WHITE = "FFFFFF"


def set_run_font(run, name="Microsoft YaHei", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_cell_text(cell, bold=False, color=INK, size=9.5):
    for p in cell.paragraphs:
        set_para_spacing(p, after=2, line=1.15)
        for r in p.runs:
            set_run_font(r, size=size, bold=bold, color=color)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    set_para_spacing(p, before=18 if level == 1 else 12, after=8 if level == 1 else 6, line=1.15)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True, color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    set_para_spacing(p, after=6, line=1.25)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run_font(r, size=10.5, bold=True, color=INK)
        r = p.add_run(text[len(bold_prefix):])
        set_run_font(r, size=10.5, color=INK)
    else:
        r = p.add_run(text)
        set_run_font(r, size=10.5, color=INK)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    set_para_spacing(p, after=4, line=1.25)
    r = p.add_run(text)
    set_run_font(r, size=10.2, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    set_para_spacing(p, after=4, line=1.25)
    r = p.add_run(text)
    set_run_font(r, size=10.2, color=INK)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, CALLOUT)
    p = cell.paragraphs[0]
    set_para_spacing(p, after=2, line=1.2)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p = cell.add_paragraph()
    set_para_spacing(p, after=0, line=1.2)
    r = p.add_run(body)
    set_run_font(r, size=10, color=INK)
    doc.add_paragraph()
    return table


def add_table(doc, headers, rows, widths, font_size=9.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        shade_cell(hdr[i], LIGHT_BLUE)
        style_cell_text(hdr[i], bold=True, color=DARK_BLUE, size=font_size)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            shade_cell(cells[i], WHITE)
            style_cell_text(cells[i], size=font_size)
    set_table_width(table, widths)
    doc.add_paragraph()
    return table


def set_doc_styles(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        st = styles[name]
        st.font.name = "Microsoft YaHei"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.line_spacing = 1.15

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Regional AI & Digital Transformation Intern Interview Prep")
    set_run_font(run, size=8.5, color=MUTED)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_para_spacing(p, before=0, after=4)
    r = p.add_run("面试资料包")
    set_run_font(r, size=12, bold=True, color=BLUE)

    p = doc.add_paragraph()
    set_para_spacing(p, after=8, line=1.05)
    r = p.add_run("Regional AI & Digital Transformation Intern")
    set_run_font(r, size=24, bold=True, color=INK)

    p = doc.add_paragraph()
    set_para_spacing(p, after=16, line=1.2)
    r = p.add_run("面向 Regional Marketing / Asia Pacific 的 AI adoption、数字化转型、营销技术和客户互动岗位准备")
    set_run_font(r, size=11, color=MUTED)

    add_callout(
        doc,
        "核心定位",
        "你要呈现的不是“会用 AI 工具”，而是“能把 AI 做成合规、可审核、可推广的营销工作流”。",
    )

    add_table(
        doc,
        ["JD 信号", "面试官真正想确认"],
        [
            ["AI Translation & Localisation", "你懂多语言内容流程、术语库、品牌语气、QA 和本地市场 review。"],
            ["AI Marketing Technology", "你能评估工具，不只追热点；会看业务价值、成本、集成、安全、ROI。"],
            ["AI Customer Engagement", "你能设计 assistant/chatbot 的知识库、边界、人工转接和失败处理。"],
            ["AI Solution Deployment", "你能从 pilot 到 rollout、hypercare、反馈迭代，而不是只做 demo。"],
            ["Documentation & Change Management", "你能写 SOP、培训用户、推动 adoption。"],
            ["Governance / Responsible AI", "你知道敏感数据、幻觉、版权、合规、人审节点和 approved platforms。"],
        ],
        [2592, 6768],
        font_size=9.3,
    )


def add_competency_map(doc):
    add_heading(doc, "1. 面试会重点考核什么", 1)
    rows = [
        ["AI 工作流设计", "能把营销任务拆成输入、AI 处理、人工审核、发布、复盘", "AI 不是一键生成，要有流程、质量门禁和责任边界"],
        ["本地化理解", "术语库、风格指南、多市场 reviewer、QA checklist", "一致性、速度、品牌安全是翻译工作流的核心指标"],
        ["MarTech 判断力", "工具评估、roadmap、优先级、ROI、数据安全", "先从业务痛点和 adoption 成本出发，再谈工具"],
        ["Customer engagement", "chatbot/assistant/RAG/知识库/升级人工", "面向客户的 AI 必须能拒答、引用来源、记录反馈"],
        ["治理与安全", "敏感数据、权限、approved platform、responsible AI", "实习生也必须懂边界，不能把安全甩给 IT"],
        ["落地推进", "pilot、UAT、training、hypercare、continuous improvement", "公司要的是能上线、能培训、能优化的人"],
        ["跨区域沟通", "APAC 语言、文化、时区、stakeholder 管理", "你需要把技术话翻译成市场团队听得懂的行动"],
    ]
    add_table(doc, ["能力", "他们怎么判断", "你要传达的观点"], rows, [1800, 3420, 4140])

    add_callout(
        doc,
        "一句话答题主线",
        "“我会先明确业务流程和风险边界，再选择 approved AI tools，把输出放进可审核的工作流，最后用 adoption、质量和效率指标证明价值。”",
    )


def add_project_packaging(doc):
    add_heading(doc, "2. 把你的 AI Web Form Agent 包装成岗位相关案例", 1)
    add_body(
        doc,
        "你的项目不要只说成 form filler。更强的说法是：review-first AI browser workflow assistant，用 AI 读取网页、提取字段、映射信息、生成可审核计划，并在人工确认后执行安全浏览器动作。",
    )
    add_table(
        doc,
        ["项目能力", "对应 JD", "面试表达"],
        [
            ["页面读取与结构化提取", "AI Solution Deployment / Customer Engagement", "我做的是从非结构化页面中提取任务信息，并转化为可执行 workflow。"],
            ["review-first / approval gates", "Governance / Responsible AI", "系统不会自动提交，关键动作前保留人审，符合 responsible AI 的落地原则。"],
            ["可复用 profile / reviewed memory", "AI Translation & Localisation / Productivity", "我理解 AI adoption 需要复用已审核知识，而不是每次重新生成。"],
            ["benchmark / trace / evidence", "Evaluation / Continuous Improvement", "我用测试、trace 和 benchmark 证明 workflow 是否稳定，而不是只靠主观体验。"],
            ["本地 demo 无 API key", "Deployment / Adoption", "我关注 demo 可复现，能让非技术面试官快速看到价值。"],
        ],
        [1900, 2400, 5060],
    )

    add_heading(doc, "2.1 两分钟项目介绍模板", 2)
    add_body(
        doc,
        "中文：我最近做了一个 AI browser workflow assistant。它不是只自动填表，而是把网页任务拆成“读取页面、提取字段、匹配用户资料、生成计划、人工审核、执行浏览器动作、验证结果”的流程。项目里我特别强调 human-in-the-loop，不自动提交表单，不处理验证码、付款或登录绕过，并且会记录截图、日志和验证 evidence。这个经验和这个岗位非常接近，因为 Regional Marketing 的 AI adoption 也需要把 AI 嵌入真实流程，比如翻译、本地化、营销助手或客户信息 chatbot，并且要有治理、QA 和 adoption 机制。",
    )
    add_body(
        doc,
        "English：One project I have been building is a review-first AI browser workflow assistant. Instead of treating automation as a one-click form filler, I designed it as a controlled workflow: read the page, extract required fields, map reviewed profile data, generate an inspectable plan, require human review, execute safe browser actions, and verify the result. I deliberately kept safety boundaries such as no auto-submission, no CAPTCHA bypass, no payment automation, and no storage of one-time sensitive values. I think this maps well to regional marketing AI adoption, where the real value is not just generating content, but building scalable, governed workflows for translation, localization, marketing support, and customer engagement.",
    )


def add_frameworks(doc):
    add_heading(doc, "3. 四个高概率 case 的答题框架", 1)
    cases = [
        (
            "Case A: 设计 APAC AI 翻译与本地化工作流",
            [
                "内容分级：campaign copy、product facts、legal/compliance、social post、internal training 分开处理。",
                "建立输入资产：brand voice guide、术语库、禁用词、产品事实库、市场差异规则。",
                "AI 初译：使用 approved internal AI platform，要求输出目标语言、语气说明、风险提示和不确定项。",
                "QA 门禁：本地市场 reviewer 检查准确性、文化适配、术语一致性、合规风险。",
                "反馈闭环：把 approved corrections 写回术语库和 prompt template。",
                "指标：turnaround time、review changes per 100 words、terminology accuracy、market satisfaction、reuse rate。",
            ],
        ),
        (
            "Case B: 评估 AI Marketing Technology",
            [
                "先定义业务问题：内容生产慢、翻译不一致、客服重复问题多、产品资料难找。",
                "评估维度：business fit、accuracy、security、integration、usability、cost、scalability、governance、vendor risk。",
                "PoC 方法：选 1-2 个真实 use cases，小范围测试，比较人工流程和 AI-assisted 流程。",
                "Roadmap：quick wins 先做内部助手和翻译 QA，再扩展到 customer-facing assistant。",
                "避免陷阱：不因为工具热门就采购；不把未经审核的输出直接面向客户。",
            ],
        ),
        (
            "Case C: 设计产品信息 chatbot / marketing assistant",
            [
                "知识源：产品手册、FAQ、campaign brief、approved claims、regional policy。",
                "RAG：检索知识库，回答时引用来源，知识过期时提示无法确认。",
                "边界：不能编造价格、医疗/法律承诺、未批准 claims、客户隐私信息。",
                "人工转接：低置信度、投诉、合规问题、购买决策等转人工。",
                "评价：answer accuracy、containment rate、handoff quality、CSAT、unsupported refusal rate。",
            ],
        ),
        (
            "Case D: 从 pilot 到 rollout",
            [
                "Discovery：访谈 regional marketing 和本地市场，找高频低风险流程。",
                "Pilot：选 1-2 个国家或一个内容类型，定义成功指标。",
                "UAT：让真实用户按 SOP 操作，记录失败案例和困惑点。",
                "Training：短视频、user manual、office hour、FAQ。",
                "Hypercare：上线后 2-4 周收集问题，快速修 prompt、模板、权限和流程。",
                "Scale：形成 playbook，再复制到更多市场和内容类型。",
            ],
        ),
    ]
    for title, points in cases:
        add_heading(doc, title, 2)
        for point in points:
            add_bullet(doc, point)


def add_question_bank(doc):
    add_heading(doc, "4. 高频面试题库：问题、考点、回答方向", 1)
    rows = [
        ["请你自我介绍。", "岗位匹配度", "AI workflow + marketing context + 项目实践 + 学习速度。"],
        ["为什么对这个岗位感兴趣？", "动机", "AI 真正价值在业务流程落地；APAC 多语言营销是高价值场景。"],
        ["你如何设计 AI 翻译流程？", "本地化能力", "术语库、brand voice、AI 初译、本地 review、QA、反馈闭环。"],
        ["如何判断一个 AI 工具值不值得用？", "MarTech 判断力", "业务价值、安全、集成、成本、准确率、adoption、ROI。"],
        ["如果 AI 输出错误怎么办？", "风险意识", "人审、来源、置信度、日志、fallback、持续改进。"],
        ["设计一个客户 chatbot。", "AI 应用设计", "知识库、RAG、边界、人工转接、指标。"],
        ["如何推动区域团队采用新 AI 工具？", "change management", "pilot、champions、training、SOP、hypercare、反馈机制。"],
        ["如何处理敏感数据？", "治理", "数据分类、approved tools、最小权限、不上传敏感信息、不存储一次性值。"],
        ["你最强的技术能力是什么？", "自我认知", "把 AI agent/workflow 做成可审核、可测试、可解释的产品。"],
        ["你不熟悉营销怎么办？", "学习能力", "承认边界，但展示学习方法：访谈、benchmark、快速 PoC、指标验证。"],
        ["你如何衡量 AI 项目成功？", "结果导向", "效率、质量、风险、adoption、用户满意、业务影响。"],
        ["你如何应对跨文化内容差异？", "APAC 协作", "本地 reviewer、市场规则、文化敏感点、统一术语和品牌 voice。"],
    ]
    add_table(doc, ["面试问题", "考点", "回答方向"], rows, [2500, 1800, 5060], font_size=8.8)

    add_heading(doc, "4.1 面试中要主动抛出的关键词", 2)
    keywords = [
        "human-in-the-loop / 人工审核",
        "approved internal AI platform",
        "brand voice guide / terminology glossary",
        "RAG / source-grounded answers",
        "quality assurance checklist",
        "data governance / responsible AI",
        "pilot, UAT, rollout, hypercare",
        "adoption metrics and continuous improvement",
        "workflow template, SOP, reusable playbook",
    ]
    for k in keywords:
        add_bullet(doc, k)


def add_answer_templates(doc):
    add_heading(doc, "5. 可背诵的参考答案", 1)
    qa = [
        (
            "Q: 你会如何建立一个 APAC AI 翻译工作流？",
            "A: 我会先把内容按风险分级，比如普通社媒文案、产品资料、合规敏感内容分别走不同审核强度。然后建立术语库、brand voice guide 和产品事实库，让 AI 初译有标准输入。AI 输出后不会直接发布，而是由本地市场 reviewer 检查语言、文化适配、产品准确性和合规风险。每次修改都沉淀回术语库和 prompt template。最后我会用 turnaround time、术语准确率、review 修改率和市场满意度衡量效果。",
        ),
        (
            "Q: 如果让你推荐一个 AI marketing technology stack，你会怎么做？",
            "A: 我不会先从工具列表开始，而会从业务优先级开始。比如区域营销最常见的痛点可能是多语言内容效率、产品信息查找、campaign asset 复用和客户问题响应。我会为每个场景定义评估标准，包括准确率、安全、集成成本、用户体验、可扩展性和 ROI。先做小范围 PoC，再决定是否进入 roadmap。初期我会优先推荐低风险高频场景，比如内部 marketing assistant、translation QA 和 campaign brief summarization。",
        ),
        (
            "Q: 面向客户的 AI assistant 最大风险是什么？",
            "A: 最大风险是它看起来很自信，但给出未经批准或错误的信息。所以我会要求 assistant 只基于 approved knowledge base 回答，重要回答给出来源；对价格、合规承诺、敏感客户数据等问题设置拒答或人工转接；对低置信度回答做 fallback。上线后要看日志、错误样本和用户反馈，不断优化知识库和边界规则。",
        ),
        (
            "Q: 你作为 intern 如何推动 adoption？",
            "A: 我会把角色定位成 enablement partner。先访谈团队，找最重复、最耗时、风险可控的流程；做一个小 pilot；把操作步骤写成 SOP 和一页 quick guide；组织短 training 和 office hour；上线后做 hypercare，收集问题并迭代模板。这样可以降低大家对 AI 的心理门槛，也能让团队看到实际节省时间和质量提升。",
        ),
    ]
    for q, a in qa:
        add_heading(doc, q, 2)
        add_body(doc, a)


def add_30_60_90(doc):
    add_heading(doc, "6. 30 / 60 / 90 天入职方案", 1)
    add_table(
        doc,
        ["阶段", "目标", "具体动作", "可交付物"],
        [
            ["前 30 天", "理解业务和治理边界", "访谈区域和本地市场；梳理内容流程；学习 approved AI platforms 和 data governance。", "痛点地图、流程现状图、风险清单、优先级 use case shortlist。"],
            ["31-60 天", "做出可验证 pilot", "选择 1-2 个低风险高频场景；设计 prompt、SOP、QA checklist；组织 UAT。", "AI 翻译/内容助手 pilot、用户手册、测试反馈、初版指标报告。"],
            ["61-90 天", "推广和沉淀", "根据反馈迭代；做 training；建立 playbook；提出 MarTech roadmap。", "rollout plan、training materials、best practices、AI Marketing Technology Stack roadmap。"],
        ],
        [1400, 1900, 3600, 2460],
        font_size=8.7,
    )


def add_interviewer_simulation(doc):
    add_heading(doc, "7. 真实考官模拟：面试官画像与追问方式", 1)
    add_body(
        doc,
        "这类岗位通常不是纯技术面试。更可能是 Regional Marketing manager、Digital transformation lead、MarTech/IT partner 或 HR 组合面试。问题会从“你理解岗位吗”逐步压到“你会怎么落地”。",
    )
    add_table(
        doc,
        ["考官角色", "会怎么问", "他们在听什么"],
        [
            ["Regional Marketing Manager", "我们 APAC 内容很多语言版本，你会如何提升效率又保证质量？", "业务流程、本地化质量、跨市场协作。"],
            ["Digital Transformation Lead", "你如何从 pilot 推到更多国家？", "项目管理、change management、可扩展方法。"],
            ["IT / Governance Partner", "你如何确保不违反数据和 AI 使用政策？", "安全意识、approved platform、人审、数据分类。"],
            ["HR", "你遇到不熟悉领域如何学习？", "成长性、沟通、抗压、动机。"],
        ],
        [2100, 3600, 3660],
        font_size=8.8,
    )

    add_heading(doc, "7.1 模拟面试流程", 2)
    flow = [
        "Warm-up: 自我介绍，为什么申请，为什么对 AI + marketing 感兴趣。",
        "JD Fit: 逐项问 translation、MarTech、customer engagement、deployment、documentation。",
        "Case Interview: 给一个 APAC 营销 AI 场景，让你现场设计 workflow。",
        "Risk Probe: 追问数据安全、AI 幻觉、合规、approved tools、客户可见内容。",
        "Behavioral: 跨团队合作、学习新工具、处理反馈、推动别人使用新流程。",
        "Candidate Questions: 你反问团队现状、AI 平台、成功指标、pilot 范围。",
    ]
    for item in flow:
        add_number(doc, item)


def add_prompt(doc):
    add_heading(doc, "8. 复制给 ChatGPT 的真实模拟面试提示词", 1)
    prompt = """你现在扮演一家跨国公司 Regional Marketing 团队的真实面试官，岗位是 Regional AI & Digital Transformation Intern，区域覆盖 Asia Pacific。请用真实、专业、有压力但不过分刁难的方式面试我。

岗位背景：
- 支持 APAC 区域营销团队加速 AI adoption 和 digital transformation。
- 工作包括 AI translation/localisation workflow、AI marketing technology assessment、AI-powered marketing assistant/chatbot、AI solution pilot/rollout/hypercare、user manual/SOP/training materials。
- 所有方案必须符合公司 data governance、security、responsible AI 和 approved internal AI platform 要求。

我的背景线索：
- 我做过一个 AI Web Form Agent / review-first AI browser workflow assistant 项目。
- 项目重点包括：读取网页、提取结构化字段、生成可审核 workflow plan、人工确认后执行、安全边界、日志/截图/verification evidence、避免自动提交和敏感操作。
- 请你帮助我把这个项目和岗位要求连接起来，但不要直接替我回答，除非我回答后请你点评。

面试规则：
1. 每次只问一个主问题，必要时追加 1-2 个追问。
2. 从 easy 到 hard：先自我介绍和动机，再问 JD 匹配，再进入 case，再问 governance 和 change management。
3. 请像真实考官一样追问细节，例如：“你具体怎么做？”“如何衡量成功？”“如果本地市场不同意怎么办？”“如果 AI 答错怎么办？”。
4. 对我的回答按 5 个维度打分：业务理解、AI workflow 设计、治理安全、落地执行、表达清晰度。每项 1-5 分。
5. 每轮点评包括：亮点、风险、可以更像强候选人的改法、一个更好的参考答案框架。
6. 请中英混合面试：大多数问题用中文，但随机穿插英文问题，尤其是自我介绍、project explanation、stakeholder communication。
7. 请覆盖这些主题：
   - AI translation and localisation workflow
   - terminology glossary and brand voice governance
   - AI marketing technology stack roadmap
   - chatbot / marketing assistant / RAG / knowledge base
   - customer-facing AI risk and hallucination control
   - approved internal AI platforms and sensitive data handling
   - pilot, UAT, rollout, hypercare, continuous improvement
   - SOP, user manual, training, adoption metrics
   - APAC cross-cultural and multilingual collaboration
   - how my AI Web Form Agent project maps to this role

请从第一题开始：让面试者做 60-90 秒自我介绍。等我回答后再继续。"""
    add_body(doc, "把下面整段复制到 ChatGPT，即可开始一轮接近真实公司的模拟面试。")
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_GRAY)
    p = cell.paragraphs[0]
    set_para_spacing(p, after=0, line=1.15)
    r = p.add_run(prompt)
    set_run_font(r, name="Consolas", size=8.6, color=INK)


def add_questions_to_ask(doc):
    add_heading(doc, "9. 你可以反问面试官的问题", 1)
    questions = [
        "目前 APAC marketing 团队最希望 AI 优先解决哪个 workflow：内容、本地化、客户互动，还是内部知识管理？",
        "公司现在有哪些 approved internal AI platforms？不同类型数据的使用边界是什么？",
        "这个 intern 角色成功的前 3 个指标会是什么？效率、adoption、质量，还是 pilot 交付？",
        "区域团队和本地市场之间，内容 review 和 localization approval 现在是怎样协作的？",
        "如果我加入，最希望我 30 天内交付的第一个小成果是什么？",
    ]
    for q in questions:
        add_bullet(doc, q)

    add_callout(
        doc,
        "最后提醒",
        "面试时避免把自己说成“只会 prompt 的人”。你要反复展示：懂业务流程、懂 AI 能力边界、懂治理、能写 SOP、能培训用户、能把小 pilot 做成可复制方案。",
    )


def main():
    doc = Document()
    set_doc_styles(doc)
    add_cover(doc)
    add_competency_map(doc)
    add_project_packaging(doc)
    add_frameworks(doc)
    add_question_bank(doc)
    add_answer_templates(doc)
    add_30_60_90(doc)
    add_interviewer_simulation(doc)
    add_prompt(doc)
    add_questions_to_ask(doc)
    doc.core_properties.title = "Regional AI & Digital Transformation Intern Interview Prep"
    doc.core_properties.subject = "Interview preparation for APAC AI marketing transformation role"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
