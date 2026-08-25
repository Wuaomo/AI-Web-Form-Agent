from pathlib import Path

from docx import Document

from create_ai_marketing_interview_doc import (
    add_body,
    add_bullet,
    add_callout,
    add_heading,
    add_number,
    add_table,
    set_doc_styles,
    set_para_spacing,
    set_run_font,
    set_table_width,
    shade_cell,
    LIGHT_GRAY,
    INK,
)


OUT = Path(r"C:\Users\wuaomo\Documents\AI Web Form Agent\Regional_AI_Digital_Transformation_Intern_English_Interview_Prep.docx")


def cover(doc):
    p = doc.add_paragraph()
    set_para_spacing(p, after=4)
    r = p.add_run("English Interview Prep Pack")
    set_run_font(r, size=12, bold=True, color="2E74B5")

    p = doc.add_paragraph()
    set_para_spacing(p, after=8, line=1.05)
    r = p.add_run("Regional AI & Digital Transformation Intern")
    set_run_font(r, size=24, bold=True, color=INK)

    add_body(
        doc,
        "A practical speaking guide for an English interview with a Regional Marketing / Asia Pacific team. The goal is to sound like someone who can turn AI tools into governed, scalable marketing workflows.",
    )

    add_callout(
        doc,
        "Core Positioning",
        "Do not present yourself as someone who only knows prompts. Present yourself as someone who understands AI workflow design, marketing operations, localisation quality, data governance, and change management.",
    )


def speaking_positioning(doc):
    add_heading(doc, "1. Your Interview Positioning", 1)
    add_table(
        doc,
        ["What the interviewer asks", "What they are really testing", "How you should position yourself"],
        [
            ["Why this role?", "Motivation and business maturity", "I am interested in using AI to improve real marketing workflows, not just experimenting with tools."],
            ["Tell me about your AI project.", "Can you connect technical work to business workflows?", "My AI Web Form Agent is a review-first browser workflow assistant with human approval and evidence logging."],
            ["How would you support APAC localisation?", "Process thinking", "I would combine terminology governance, brand voice guidance, AI-assisted translation, local market review, and QA metrics."],
            ["How do you evaluate AI tools?", "Judgment", "I would assess business fit, accuracy, security, integration effort, adoption friction, scalability, and ROI."],
            ["What about AI risks?", "Governance awareness", "I would keep sensitive data out, use approved platforms, require human review, and define clear answer boundaries."],
        ],
        [2300, 3100, 3960],
        font_size=8.8,
    )

    add_callout(
        doc,
        "One-Sentence Thesis",
        "My strength is connecting AI capability with practical marketing workflows, so the solution is useful, reviewable, safe, and easy for regional teams to adopt.",
    )


def self_intro(doc):
    add_heading(doc, "2. 60-90 Second Self-Introduction", 1)
    add_body(
        doc,
        "Use this as your base version. Speak naturally; do not memorize it word for word.",
    )
    add_callout(
        doc,
        "Sample Answer",
        "Hi, my name is [Name]. I am interested in the intersection of AI, marketing operations, and digital transformation. Recently, I have been building an AI Web Form Agent, which I position as a review-first AI browser workflow assistant. The system reads a web page, extracts structured fields, maps information, creates an inspectable workflow plan, and only executes after human review. I paid special attention to safety boundaries, such as not auto-submitting forms, not bypassing login or CAPTCHA, and not storing sensitive one-time data. What excites me about this internship is that the role is not just about using AI tools, but about applying AI responsibly to real regional marketing workflows, including localisation, marketing assistants, customer engagement, documentation, and adoption across Asia Pacific. I believe my project experience and my learning mindset can help me contribute quickly while continuing to grow in a business environment.",
    )

    add_heading(doc, "2.1 Shorter Version", 2)
    add_body(
        doc,
        "I am interested in building AI workflows that are practical, safe, and easy for business users to adopt. My recent AI Web Form Agent project helped me understand how to design human-in-the-loop automation, define safety boundaries, and verify outputs with evidence. I would like to apply the same mindset to marketing workflows such as localisation, AI assistants, and regional digital transformation.",
    )


def project_pitch(doc):
    add_heading(doc, "3. How to Explain Your AI Web Form Agent Project", 1)
    add_body(
        doc,
        "The strongest version is not 'I built a form filler.' Say this instead:",
    )
    add_callout(
        doc,
        "Project Pitch",
        "One project I have been building is a review-first AI browser workflow assistant. Instead of treating automation as a one-click form filler, I designed it as a controlled workflow: read the page, extract required fields, map reviewed profile data, generate an inspectable plan, require human review, execute safe browser actions, and verify the result. This experience is relevant to regional marketing AI adoption because many marketing use cases, such as translation, localisation QA, internal marketing assistants, and customer-facing chatbots, also require structured workflows, human review, compliance boundaries, and measurable improvement.",
    )

    add_table(
        doc,
        ["Project Feature", "Business Meaning", "Role Connection"],
        [
            ["Page reading and field extraction", "Turns messy web tasks into structured workflows", "Useful for marketing workflow automation and digital experience improvement"],
            ["Human review before execution", "Prevents risky AI automation", "Matches responsible AI and governance requirements"],
            ["Policy gates", "Blocks unsafe actions", "Relevant to approved platform and data handling standards"],
            ["Logs, screenshots, verification evidence", "Makes automation auditable", "Supports rollout, hypercare, and continuous improvement"],
            ["Local demo without LLM keys", "Easy to reproduce and explain", "Shows practical product thinking, not just model dependency"],
        ],
        [2200, 3300, 3860],
        font_size=8.8,
    )


def case_frameworks(doc):
    add_heading(doc, "4. Case Answer Frameworks", 1)

    add_heading(doc, "Case 1: Design an AI Translation and Localisation Workflow", 2)
    for item in [
        "First, I would classify content by risk level: social copy, campaign copy, product information, legal or compliance-sensitive content.",
        "Then I would create the core assets: terminology glossary, brand voice guide, approved product claims, and local market rules.",
        "The AI platform would generate a first draft, but the output would go through local market review before publication.",
        "The QA checklist would cover terminology accuracy, brand tone, factual accuracy, cultural sensitivity, and compliance issues.",
        "Approved corrections should feed back into the glossary and prompt templates.",
        "Success metrics: turnaround time, review change rate, terminology accuracy, reuse rate, and local market satisfaction.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Case 2: Evaluate an AI Marketing Technology Stack", 2)
    for item in [
        "I would start from business priorities, not from a tool list.",
        "Typical APAC marketing pain points could include multilingual content production, asset reuse, product information access, and repeated customer questions.",
        "My evaluation criteria would include business fit, output quality, data security, integration effort, user experience, scalability, cost, vendor risk, and governance support.",
        "I would run a small proof of concept using real but low-risk workflows, then compare AI-assisted performance with the current manual process.",
        "The roadmap should start with low-risk internal use cases before moving into customer-facing AI.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Case 3: Design a Customer-Facing AI Assistant", 2)
    for item in [
        "The assistant should answer only from approved sources such as product FAQs, campaign briefs, product manuals, and approved claims.",
        "I would use retrieval-augmented generation so the assistant can ground answers in source documents.",
        "It should refuse or escalate questions about pricing, legal claims, medical claims, sensitive personal data, or unsupported information.",
        "Low-confidence answers should trigger a fallback or human handoff.",
        "Key metrics: answer accuracy, source coverage, unsupported refusal rate, handoff quality, customer satisfaction, and issue resolution rate.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Case 4: Pilot and Rollout Plan", 2)
    for item in [
        "Discovery: interview regional and local market stakeholders to identify high-frequency, low-risk workflows.",
        "Pilot: test one or two use cases in a small market group with clear success metrics.",
        "UAT: let real users test the workflow and capture friction points.",
        "Training: provide user manuals, SOPs, short demos, office hours, and FAQ materials.",
        "Hypercare: monitor issues for the first few weeks, refine prompts, templates, workflows, and governance rules.",
        "Scale: convert the pilot into a reusable playbook for additional markets.",
    ]:
        add_bullet(doc, item)


def question_bank(doc):
    add_heading(doc, "5. High-Probability English Interview Questions", 1)
    add_table(
        doc,
        ["Question", "What to Show", "Answer Direction"],
        [
            ["Tell me about yourself.", "Clear positioning", "AI workflow, marketing operations, project experience, learning mindset."],
            ["Why are you interested in this role?", "Motivation", "AI adoption in real business workflows, APAC localisation complexity."],
            ["How would you design an AI translation workflow?", "Process thinking", "Glossary, brand voice, AI draft, local review, QA, feedback loop."],
            ["How would you evaluate AI marketing tools?", "Business judgment", "Business fit, accuracy, governance, integration, cost, adoption, ROI."],
            ["How would you prevent AI hallucination?", "Risk awareness", "Approved sources, citations, confidence thresholds, human review, escalation."],
            ["How would you build a chatbot for product information?", "Solution design", "Knowledge base, RAG, answer boundaries, handoff, metrics."],
            ["How would you drive adoption across regional teams?", "Change management", "Pilot, champions, training, SOPs, hypercare, feedback loop."],
            ["How do you handle sensitive data?", "Governance", "Data classification, approved tools, least privilege, no sensitive data in prompts."],
            ["What is your weakness for this role?", "Self-awareness", "Be honest: marketing domain depth, then show structured learning plan."],
            ["What would you deliver in your first 90 days?", "Execution", "Pain point map, pilot, SOP/training, roadmap, measurable results."],
        ],
        [2900, 2100, 4360],
        font_size=8.7,
    )


def polished_answers(doc):
    add_heading(doc, "6. Polished Sample Answers", 1)
    samples = [
        (
            "Q: How would you design an AI-enabled translation workflow for APAC?",
            "I would start by separating content into different risk levels, because a social post, a product claim, and a compliance-sensitive document should not follow the same process. Then I would build the reusable assets: a terminology glossary, brand voice guide, approved product claims, and local market rules. The AI tool would create the first draft, but local market reviewers would validate accuracy, tone, cultural fit, and compliance. I would also create a QA checklist and feed approved corrections back into the glossary and prompt templates. The success metrics would include turnaround time, terminology accuracy, review change rate, reuse rate, and satisfaction from local teams.",
        ),
        (
            "Q: How would you evaluate AI marketing technologies?",
            "I would not start by asking which tool is the most advanced. I would start by identifying the business problem we want to solve, such as slow localisation, inconsistent campaign content, or repeated customer questions. Then I would assess each tool against business fit, output quality, data security, integration effort, user experience, scalability, cost, and governance support. I would recommend a small proof of concept first, using real but low-risk workflows, and compare the AI-assisted process with the current manual process before adding it to the roadmap.",
        ),
        (
            "Q: What are the main risks of a customer-facing AI assistant?",
            "The biggest risk is that the assistant may sound confident while giving inaccurate or unapproved information. To manage this, I would restrict it to approved knowledge sources, use source-grounded answers, and define clear boundaries for what it can and cannot answer. For sensitive, legal, pricing, or low-confidence questions, it should escalate to a human or provide a safe fallback. I would also review logs and failure cases after launch so the knowledge base and policies can improve continuously.",
        ),
        (
            "Q: How would you encourage regional marketing teams to adopt a new AI workflow?",
            "I would treat adoption as a change management problem, not just a technology rollout. First, I would understand users' current workflow and pain points. Then I would start with a small pilot that saves time in a visible way. I would provide a simple SOP, a short training session, and examples of good outputs. After launch, I would run hypercare sessions, collect feedback, and improve the workflow. The goal is to make AI feel practical and safe, rather than forcing people to use a tool they do not trust.",
        ),
    ]
    for q, a in samples:
        add_heading(doc, q, 2)
        add_body(doc, a)


def ninety_day_plan(doc):
    add_heading(doc, "7. 30 / 60 / 90 Day Plan", 1)
    add_table(
        doc,
        ["Period", "Goal", "Actions", "Deliverables"],
        [
            ["First 30 days", "Understand workflows and governance", "Interview stakeholders, map localisation and marketing workflows, learn approved AI platforms and data rules.", "Pain point map, current process map, risk checklist, prioritised use case list."],
            ["Days 31-60", "Build and test a pilot", "Select one or two low-risk use cases, design prompts, SOPs, QA checklist, and run UAT.", "Pilot workflow, user guide, UAT feedback, first metrics report."],
            ["Days 61-90", "Roll out and improve", "Refine the workflow, train users, create reusable playbooks, propose a MarTech roadmap.", "Rollout plan, training materials, best practice guide, AI Marketing Technology Stack roadmap."],
        ],
        [1500, 2100, 3600, 2160],
        font_size=8.5,
    )


def simulator_prompt(doc):
    add_heading(doc, "8. Copy-Paste Prompt for English Mock Interview", 1)
    prompt = """You are now acting as a realistic interviewer from a multinational company's Regional Marketing team. The role is Regional AI & Digital Transformation Intern, supporting AI adoption and digital transformation across Asia Pacific.

Interview context:
- The role supports APAC marketing teams with AI translation and localisation workflows, AI marketing technology assessment, AI-powered marketing assistants or chatbots, AI solution pilots, rollout, hypercare, SOPs, user manuals, and training.
- Every proposed solution must respect data governance, security, responsible AI standards, and approved internal AI platforms.

Candidate context:
- I have built an AI Web Form Agent, which I want to position as a review-first AI browser workflow assistant.
- The project includes page reading, structured field extraction, inspectable workflow planning, human review before execution, safety boundaries, logs, screenshots, and verification evidence.
- Please help me connect this project to the role, but do not answer for me unless I ask for feedback after my answer.

Interview rules:
1. Conduct the entire interview in English.
2. Ask one main question at a time, then ask one or two realistic follow-up questions.
3. Start with self-introduction and motivation, then move into role fit, case questions, AI governance, stakeholder management, and change management.
4. Make the interview realistic: ask questions like "How exactly would you do that?", "How would you measure success?", "What if a local market disagrees?", "What if the AI gives a wrong answer?", and "How would you handle sensitive data?"
5. After each answer, score me from 1 to 5 on business understanding, AI workflow design, governance and safety, execution ability, and communication clarity.
6. After each answer, give me concise feedback: strengths, weaknesses, a stronger structure, and a polished sample answer.
7. Cover these topics during the interview:
   - AI translation and localisation workflow
   - terminology glossary and brand voice governance
   - AI marketing technology stack roadmap
   - chatbot, marketing assistant, RAG, and knowledge base design
   - customer-facing AI risks and hallucination control
   - approved internal AI platforms and sensitive data handling
   - pilot, UAT, rollout, hypercare, and continuous improvement
   - SOPs, user manuals, training, and adoption metrics
   - APAC multilingual and cross-cultural collaboration
   - how my AI Web Form Agent project maps to this role

Begin with the first question: ask me to give a 60-90 second self-introduction in English. Wait for my answer before continuing."""
    add_body(doc, "Use this prompt when you want ChatGPT to simulate the English interview.")
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, LIGHT_GRAY)
    r = cell.paragraphs[0].add_run(prompt)
    set_run_font(r, name="Consolas", size=8.4, color=INK)


def questions_to_ask(doc):
    add_heading(doc, "9. Smart Questions to Ask the Interviewer", 1)
    for q in [
        "Which marketing workflows are currently the highest priority for AI adoption in APAC?",
        "What approved internal AI platforms does the team currently use?",
        "How does the team currently manage localisation review between regional and local markets?",
        "What would success look like for this intern in the first 90 days?",
        "Is the team more focused on internal productivity use cases or customer-facing AI experiences at this stage?",
    ]:
        add_bullet(doc, q)

    add_callout(
        doc,
        "Final Reminder",
        "In English, keep answers structured: context, approach, risk control, metric. That structure will make you sound calmer and more senior.",
    )


def main():
    doc = Document()
    set_doc_styles(doc)
    cover(doc)
    speaking_positioning(doc)
    self_intro(doc)
    project_pitch(doc)
    case_frameworks(doc)
    question_bank(doc)
    polished_answers(doc)
    ninety_day_plan(doc)
    simulator_prompt(doc)
    questions_to_ask(doc)
    doc.core_properties.title = "Regional AI & Digital Transformation Intern English Interview Prep"
    doc.core_properties.subject = "English interview preparation"
    doc.core_properties.author = "Codex"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
